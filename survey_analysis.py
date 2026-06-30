# -*- coding: utf-8 -*-
"""
Raya University - Liaison Office Guest House
Guest Satisfaction Survey Analysis 2018
"""

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# ============================================================
# RAYA UNIVERSITY BRANDING
# ============================================================

UNIVERSITY_NAME = "Raya University"
DEPARTMENT = "Liaison Office Guest House"
SURVEY_YEAR = "2018"

# Raya University Colors
RAYA_NAVY = "#1A2A5E"
RAYA_GOLD = "#C89B3C"
RAYA_LIGHT_BLUE = "#4A7FC1"
RAYA_CREAM = "#F5F0E6"
RAYA_DARK = "#0D1A3A"

# Color palette for charts
NAVY = RAYA_NAVY
BLUE = RAYA_LIGHT_BLUE
LIGHT_BLUE = "#9DC3E6"
GREEN = "#63BE7B"
LIGHT_GREEN = "#A9D18E"
YELLOW = "#FFEB84"
ORANGE = "#F4B183"
RED = "#E06666"
GREY = "#595959"
DARK_GREY = "#333333"
GOLD = RAYA_GOLD

plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['font.size'] = 10
plt.rcParams['figure.dpi'] = 150
plt.rcParams['axes.edgecolor'] = RAYA_NAVY
plt.rcParams['axes.linewidth'] = 0.8

# ============================================================
# MAPPINGS
# ============================================================

SATISFACTION_MAP = {'ሀ': 5, 'ለ': 4, 'ሐ': 3, 'መ': 2, 'ሠ': 1, 'ሰ': 1}
SATISFACTION_LABEL_MAP = {
    'ሀ': 'Very Satisfied', 'ለ': 'Satisfied', 'ሐ': 'Neutral',
    'መ': 'Dissatisfied', 'ሠ': 'Very Dissatisfied', 'ሰ': 'Very Dissatisfied'
}

QUALITY_MAP = {'ሀ': 1, 'ለ': 2, 'ሐ': 3, 'መ': 4, 'ሠ': 5, 'ሰ': 5}
QUALITY_LABEL_MAP = {
    'ሀ': 'Very Low', 'ለ': 'Low', 'ሐ': 'Medium', 'መ': 'Good',
    'ሠ': 'Very Good', 'ሰ': 'Very Good'
}

PURPOSE_MAP = {'ሀ': 'Work', 'ለ': 'Training', 'ሐ': 'Meeting', 'መ': 'Other'}
DURATION_MAP = {'ሀ': '1 day', 'ለ': '2-3 days', 'ሐ': 'More than 3 days'}
YES_NO_MAP = {'ሀ': 'Yes', 'ለ': 'No', 'ሐ': 'Maybe'}

COLUMN_NAMES = [
    'respondent_id', 'visit_purpose', 'stay_duration',
    'q1_reception_speed', 'q2_staff_respect', 'q3_info_clarity', 'q4_hospitality',
    'q5_room_cleanliness', 'q6_bed_quality', 'q7_water_bathroom', 'q8_electricity_basic',
    'q9_environment_safety', 'q10_cleaning_service', 'q11_wifi_communication',
    'q12_problem_resolution', 'q13_overall_service', 'q14_return_willingness',
    'q15_recommend', 'q16_liaison_staff_capacity', 'open_1', 'open_2', 'open_3'
]

QUESTION_LABELS = {
    'q1_reception_speed': 'Reception Speed',
    'q2_staff_respect': 'Staff Respect & Courtesy',
    'q3_info_clarity': 'Information Clarity',
    'q4_hospitality': 'Hospitality',
    'q5_room_cleanliness': 'Room Cleanliness',
    'q6_bed_quality': 'Bed & Furniture Quality',
    'q7_water_bathroom': 'Water & Bathroom Service',
    'q8_electricity_basic': 'Electricity & Basic Services',
    'q9_environment_safety': 'Environmental Safety & Security',
    'q10_cleaning_service': 'Cleaning Service',
    'q11_wifi_communication': 'Wi-Fi & Communication',
    'q12_problem_resolution': 'Problem Resolution Speed',
    'q13_overall_service': 'Overall Service',
    'q16_liaison_staff_capacity': 'Liaison Staff Capacity'
}

# ============================================================
# DATA LOADING & PROCESSING
# ============================================================

def load_data():
    """Load data from Excel file"""
    data_file = None
    for f in os.listdir('.'):
        if f.endswith('.xlsx') and ('መልስ' in f or '2018' in f or 'data' in f):
            data_file = f
            break
    
    if not data_file:
        print("=" * 60)
        print("ERROR: Could not find the Excel data file!")
        print("=" * 60)
        print("\nPlease make sure your Excel file is in this folder.")
        print("Looking for file containing 'መልስ', '2018', or 'data'")
        print("\nFiles found in current folder:")
        for f in os.listdir('.'):
            print(f"  - {f}")
        return None
    
    print(f"✅ Found data file: {data_file}")
    print(f"Loading data...")
    df = pd.read_excel(data_file, sheet_name='Sheet1', header=None, skiprows=2)
    n_cols = min(len(COLUMN_NAMES), len(df.columns))
    df.columns = COLUMN_NAMES[:n_cols]
    df = df.dropna(how='all').reset_index(drop=True)
    
    for col in df.select_dtypes(include=['object']).columns:
        if col in df.columns:
            df[col] = df[col].astype(str).str.strip().replace('nan', np.nan).replace('', np.nan)
    
    df = df[pd.to_numeric(df['respondent_id'], errors='coerce').notna()].reset_index(drop=True)
    print(f"✅ Loaded {len(df)} respondents")
    return df

def process_data(df):
    """Process the survey data with CORRECT mappings"""
    print("Processing data...")
    
    def map_response(value, mapping):
        if pd.isna(value) or value == '' or value == 'nan':
            return np.nan
        val_str = str(value)
        if val_str and len(val_str) > 0:
            return mapping.get(val_str[0], np.nan)
        return np.nan
    
    sat_questions = [
        'q1_reception_speed', 'q2_staff_respect', 'q3_info_clarity', 'q4_hospitality',
        'q5_room_cleanliness', 'q6_bed_quality', 'q7_water_bathroom', 'q8_electricity_basic',
        'q9_environment_safety', 'q10_cleaning_service', 'q11_wifi_communication',
        'q12_problem_resolution'
    ]
    
    for col in sat_questions:
        if col in df.columns:
            df[f'{col}_num'] = df[col].apply(lambda x: map_response(x, SATISFACTION_MAP))
            df[f'{col}_label'] = df[col].apply(lambda x: map_response(x, SATISFACTION_LABEL_MAP))
    
    if 'q13_overall_service' in df.columns:
        df['q13_overall_service_num'] = df['q13_overall_service'].apply(lambda x: map_response(x, QUALITY_MAP))
        df['q13_overall_service_label'] = df['q13_overall_service'].apply(lambda x: map_response(x, QUALITY_LABEL_MAP))
    
    if 'q16_liaison_staff_capacity' in df.columns:
        df['q16_liaison_staff_capacity_num'] = df['q16_liaison_staff_capacity'].apply(lambda x: map_response(x, QUALITY_MAP))
        df['q16_liaison_staff_capacity_label'] = df['q16_liaison_staff_capacity'].apply(lambda x: map_response(x, QUALITY_LABEL_MAP))
    
    if 'visit_purpose' in df.columns:
        df['visit_purpose_label'] = df['visit_purpose'].apply(lambda x: map_response(x, PURPOSE_MAP))
    if 'stay_duration' in df.columns:
        df['stay_duration_label'] = df['stay_duration'].apply(lambda x: map_response(x, DURATION_MAP))
    if 'q14_return_willingness' in df.columns:
        df['q14_return_willingness_label'] = df['q14_return_willingness'].apply(lambda x: map_response(x, YES_NO_MAP))
    if 'q15_recommend' in df.columns:
        df['q15_recommend_label'] = df['q15_recommend'].apply(lambda x: 
            {'ሀ': 'Yes', 'ለ': 'No'}.get(str(x)[0] if pd.notna(x) else '', np.nan))
    
    print(f"✅ Processed {len(df)} records")
    return df

# ============================================================
# CHART GENERATION
# ============================================================

def create_charts(df):
    """Generate all visualization charts"""
    os.makedirs('charts', exist_ok=True)
    color_seq = [GREEN, LIGHT_GREEN, YELLOW, ORANGE, RED]
    
    # Chart 1: Mean Scores
    print("  Creating chart_means.png...")
    means = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        m = df[col].mean()
        if not pd.isna(m):
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            means.append({'variable': label, 'mean': m})
    
    if means:
        df_m = pd.DataFrame(means).sort_values('mean', ascending=True)
        fig, ax = plt.subplots(figsize=(10, 7))
        colors = [GREEN if m >= 4.5 else LIGHT_GREEN if m >= 4.0 else ORANGE if m >= 3.5 else RED for m in df_m['mean']]
        bars = ax.barh(df_m['variable'], df_m['mean'], color=colors, height=0.6, edgecolor='white', linewidth=0.5)
        ax.set_xlim(1, 5)
        ax.set_xlabel('Mean Score (1 = Worst → 5 = Best)', fontsize=11, color=GREY)
        ax.set_title(f'{UNIVERSITY_NAME}\nMean Satisfaction Scores by Service Area', 
                    fontsize=14, fontweight='bold', color=RAYA_NAVY, pad=15)
        ax.axvline(3, color='#999999', linestyle='--', linewidth=1)
        ax.text(3, -0.5, 'neutral', fontsize=9, color='#888888', ha='center')
        for bar, m in zip(bars, df_m['mean']):
            ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2, f'{m:.2f}', va='center', fontsize=9)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.tick_params(axis='y', labelsize=9)
        plt.tight_layout()
        plt.savefig('charts/chart_means.png', dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
    
    # Chart 2: Visit Profile
    print("  Creating chart_visit_profile.png...")
    fig, axes = plt.subplots(1, 2, figsize=(10, 4))
    fig.suptitle(f'{UNIVERSITY_NAME} - Guest Profile', fontsize=12, fontweight='bold', color=RAYA_NAVY)
    
    if 'visit_purpose_label' in df.columns:
        freq = df['visit_purpose_label'].value_counts()
        if len(freq) > 0:
            axes[0].bar(freq.index, freq.values, color=RAYA_NAVY, width=0.5)
            axes[0].set_title('Purpose of Visit', fontsize=12, fontweight='bold', color=RAYA_NAVY)
            for i, (label, count) in enumerate(freq.items()):
                if not pd.isna(label):
                    axes[0].text(i, count + 0.3, str(count), ha='center', fontsize=10, fontweight='bold')
            axes[0].spines['top'].set_visible(False)
            axes[0].spines['right'].set_visible(False)
            axes[0].set_ylabel('Number of Respondents', fontsize=10)
            axes[0].tick_params(axis='x', rotation=15, labelsize=9)
    
    if 'stay_duration_label' in df.columns:
        freq = df['stay_duration_label'].value_counts()
        if len(freq) > 0:
            axes[1].bar(freq.index, freq.values, color=RAYA_GOLD, width=0.5)
            axes[1].set_title('Length of Stay', fontsize=12, fontweight='bold', color=RAYA_NAVY)
            for i, (label, count) in enumerate(freq.items()):
                if not pd.isna(label):
                    axes[1].text(i, count + 0.3, str(count), ha='center', fontsize=10, fontweight='bold')
            axes[1].spines['top'].set_visible(False)
            axes[1].spines['right'].set_visible(False)
            axes[1].tick_params(axis='x', labelsize=9)
    
    plt.tight_layout()
    plt.savefig('charts/chart_visit_profile.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    
    # Chart 3: Loyalty Indicators
    print("  Creating chart_loyalty.png...")
    fig, ax = plt.subplots(figsize=(8, 4.5))
    
    r_dict = {'Yes': 0, 'No': 0, 'Maybe': 0}
    if 'q14_return_willingness_label' in df.columns:
        for label, count in df['q14_return_willingness_label'].value_counts().items():
            if not pd.isna(label) and label in r_dict:
                r_dict[label] = count
    
    rec_dict = {'Yes': 0, 'No': 0, 'Maybe': 0}
    if 'q15_recommend_label' in df.columns:
        for label, count in df['q15_recommend_label'].value_counts().items():
            if not pd.isna(label) and label in rec_dict:
                rec_dict[label] = count
    
    total_r = sum(r_dict.values()) or 1
    total_rec = sum(rec_dict.values()) or 1
    cats = ['Yes', 'No', 'Maybe']
    ret_vals = [r_dict[c]/total_r*100 for c in cats]
    rec_vals = [rec_dict[c]/total_rec*100 for c in cats]
    
    x = range(len(cats))
    width = 0.35
    bars1 = ax.bar([i - width/2 for i in x], ret_vals, width=width, label='Willing to Return', color=RAYA_NAVY)
    bars2 = ax.bar([i + width/2 for i in x], rec_vals, width=width, label='Would Recommend', color=RAYA_GOLD)
    
    ax.set_xticks(list(x))
    ax.set_xticklabels(cats, fontsize=10)
    ax.set_ylabel('% of Valid Responses', fontsize=10)
    ax.set_title(f'{UNIVERSITY_NAME}\nLoyalty Indicators', fontsize=13, fontweight='bold', color=RAYA_NAVY)
    
    for bars in (bars1, bars2):
        for bar in bars:
            h = bar.get_height()
            if h > 0:
                ax.text(bar.get_x() + bar.get_width()/2, h + 1.5, f'{h:.0f}%', ha='center', fontsize=9, fontweight='bold')
    
    ax.legend(frameon=False, fontsize=10)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_ylim(0, 110)
    
    plt.tight_layout()
    plt.savefig('charts/chart_loyalty.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    
    # Chart 4: Distribution
    print("  Creating chart_distribution.png...")
    focus_vars = ['q7_water_bathroom_num', 'q13_overall_service_num', 'q16_liaison_staff_capacity_num']
    focus_labels = ['Water & Bathroom Service', 'Overall Service', 'Liaison Staff Capacity']
    fig, axes = plt.subplots(len(focus_vars), 1, figsize=(9, 5.5))
    order = ['Very Good', 'Good', 'Medium', 'Low', 'Very Low']
    
    for idx, (ax, var, label) in enumerate(zip(axes, focus_vars, focus_labels)):
        if var in df.columns:
            freq = df[var].value_counts()
            vals = [freq.get(v, 0) for v in [5, 4, 3, 2, 1]]
            total = sum(vals) or 1
            pct_vals = [v/total*100 for v in vals]
            left = 0
            for val, pct, col in zip(order, pct_vals, color_seq):
                ax.barh([label], [pct], left=left, color=col, height=0.5, edgecolor='white', linewidth=0.5)
                if pct >= 3:
                    ax.text(left + pct/2, 0, f'{pct:.0f}%', ha='center', va='center', fontsize=8,
                           fontweight='bold', color='white' if pct > 30 else DARK_GREY)
                left += pct
            ax.set_xlim(0, 100)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_visible(False)
            ax.tick_params(axis='y', labelsize=10)
            ax.set_yticks([0])
            ax.set_yticklabels([label], fontsize=10, fontweight='bold')
    
    axes[-1].set_xlabel('% of Valid Responses', fontsize=10, color=GREY)
    fig.suptitle(f'{UNIVERSITY_NAME}\nDistribution of Ratings for Key Service Areas', 
                fontsize=13, fontweight='bold', color=RAYA_NAVY, y=0.98)
    
    legend_items = [Patch(facecolor=c, label=o) for c, o in zip(color_seq, order)]
    fig.legend(handles=legend_items, loc='lower center', ncol=5, frameon=False, fontsize=8, bbox_to_anchor=(0.5, -0.02))
    
    plt.tight_layout(rect=[0, 0.06, 1, 0.94])
    plt.savefig('charts/chart_distribution.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    
    print("  ✅ All charts created!")

# ============================================================
# EXCEL EXPORT
# ============================================================

def export_excel(df):
    """Export all analysis results to Excel"""
    print("  Exporting survey_analysis_2018.xlsx...")
    with pd.ExcelWriter('survey_analysis_2018.xlsx', engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Raw Data', index=False)
        
        stats = []
        for col in [c for c in df.columns if c.endswith('_num')]:
            data = df[col].dropna()
            if len(data) > 0:
                label = QUESTION_LABELS.get(col.replace('_num', ''), col)
                stats.append({
                    'Service Area': label,
                    'N': len(data),
                    'Mean': round(data.mean(), 2),
                    'Median': round(data.median(), 2),
                    'Std Dev': round(data.std(), 2),
                    'Min': data.min(),
                    'Max': data.max()
                })
        if stats:
            pd.DataFrame(stats).to_excel(writer, sheet_name='Summary Stats', index=False)
        
        for col in [c for c in df.columns if c.endswith('_label')]:
            if col not in ['q14_return_willingness_label', 'q15_recommend_label']:
                freq = df[col].value_counts(dropna=False)
                valid = df[col].dropna().value_counts()
                total = len(df)
                valid_total = len(df[col].dropna()) or 1
                freq_df = pd.DataFrame({
                    'Response': freq.index,
                    'Count': freq.values,
                    'Percent': [round(v/total*100, 1) for v in freq.values],
                    'Valid %': [round(v/valid_total*100, 1) for v in 
                               [valid.get(x, 0) if not pd.isna(x) else 0 for x in freq.index]]
                })
                base = col.replace('_label', '')
                sheet = QUESTION_LABELS.get(base, base)[:31]
                freq_df.to_excel(writer, sheet_name=sheet, index=False)
        
        for col in ['visit_purpose_label', 'stay_duration_label']:
            if col in df.columns:
                freq = df[col].value_counts()
                pd.DataFrame({
                    'Category': freq.index,
                    'Count': freq.values,
                    'Percent': [round(v/len(df)*100, 1) for v in freq.values]
                }).to_excel(writer, sheet_name=col.replace('_label', '')[:31], index=False)
        
        loyalty = []
        for col in ['q14_return_willingness_label', 'q15_recommend_label']:
            if col in df.columns:
                for val, count in df[col].value_counts(dropna=False).items():
                    loyalty.append({
                        'Indicator': col.replace('_label', '').replace('_', ' ').title(),
                        'Response': val if not pd.isna(val) else 'Missing',
                        'Count': count,
                        'Percent': round(count/len(df)*100, 1)
                    })
        if loyalty:
            pd.DataFrame(loyalty).to_excel(writer, sheet_name='Loyalty', index=False)
    
    print("  ✅ Excel export complete!")

# ============================================================
# WORD REPORT
# ============================================================

def generate_word_report(df):
    """Generate Word document report"""
    print("  Generating survey_report_2018.docx...")
    doc = Document()
    
    title = doc.add_heading(f'{UNIVERSITY_NAME}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f'{DEPARTMENT}\nGuest Satisfaction Survey Report {SURVEY_YEAR}', 1)
    doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    n = len(df)
    all_means = [df[col].mean() for col in df.columns if col.endswith('_num') and not pd.isna(df[col].mean())]
    overall_mean = np.mean(all_means) if all_means else 0
    
    means_data = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        m = df[col].mean()
        if not pd.isna(m):
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            means_data.append({'label': label, 'mean': m})
    means_df = pd.DataFrame(means_data).sort_values('mean', ascending=False)
    top_3 = means_df.head(3)['label'].tolist()
    
    return_rate = round(df['q14_return_willingness_label'].value_counts().get('Yes', 0) / len(df['q14_return_willingness_label'].dropna()) * 100, 1) if 'q14_return_willingness_label' in df.columns else 0
    rec_rate = round(df['q15_recommend_label'].value_counts().get('Yes', 0) / len(df['q15_recommend_label'].dropna()) * 100, 1) if 'q15_recommend_label' in df.columns else 0
    
    exec_summary = f"""
A guest satisfaction survey was conducted at the {DEPARTMENT} to assess the quality of service provided to visitors. A total of {n} guests completed the survey, rating their experience across reception, accommodation, support services, and overall satisfaction.

Overall, guests reported a high level of satisfaction, with an average score of {overall_mean:.2f} out of 5 across all measured service dimensions. The strongest-performing areas were {top_3[0] if len(top_3)>0 else 'N/A'}, {top_3[1] if len(top_3)>1 else 'N/A'}, and {top_3[2] if len(top_3)>2 else 'N/A'}, while water and bathroom service received the most varied feedback and represents the clearest opportunity for improvement.

{return_rate}% of respondents indicated they would return to the guesthouse, and {rec_rate}% said they would recommend it to others — both strong indicators of guest loyalty and satisfaction.
"""
    doc.add_paragraph(exec_summary.strip())
    doc.add_page_break()
    
    # Section 2: Methodology
    doc.add_heading('2. Methodology', level=1)
    methodology = f"""
Data was collected through a structured, bilingual (Amharic/English) paper-based questionnaire administered to guests of the {DEPARTMENT}. Responses were later digitized and analyzed.

• Sample size: {n} respondents
• Survey instrument: 18-item structured questionnaire across six sections
• Rating scale: 5-point Likert scale (1 = Very dissatisfied/Very low, 5 = Very satisfied/Very good)
• Data collection method: Manual paper survey, later digitized for analysis
"""
    doc.add_paragraph(methodology.strip())
    doc.add_page_break()
    
    # Section 3: Respondent Profile
    doc.add_heading('3. Respondent Profile', level=1)
    if 'visit_purpose_label' in df.columns:
        doc.add_heading('3.1 Purpose of Visit', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Purpose'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
        for purpose, count in df['visit_purpose_label'].value_counts().items():
            if not pd.isna(purpose):
                row = table.add_row().cells
                row[0].text = str(purpose)
                row[1].text = str(count)
                row[2].text = f'{round(count/len(df)*100, 1)}%'
    
    if 'stay_duration_label' in df.columns:
        doc.add_heading('3.2 Length of Stay', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Duration'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
        for duration, count in df['stay_duration_label'].value_counts().items():
            if not pd.isna(duration):
                row = table.add_row().cells
                row[0].text = str(duration)
                row[1].text = str(count)
                row[2].text = f'{round(count/len(df)*100, 1)}%'
    doc.add_page_break()
    
    # Section 4: Overall Satisfaction
    doc.add_heading('4. Overall Satisfaction', level=1)
    if 'q13_overall_service_num' in df.columns:
        overall_score = df['q13_overall_service_num'].mean()
        doc.add_paragraph(f"The overall service satisfaction score averages {overall_score:.2f} out of 5.0.")
        doc.add_heading('4.1 Distribution of Overall Service Ratings', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Rating'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
        for rating in ['Very Good', 'Good', 'Medium', 'Low', 'Very Low']:
            count = df['q13_overall_service_label'].value_counts().get(rating, 0)
            if count > 0:
                row = table.add_row().cells
                row[0].text = rating
                row[1].text = str(count)
                row[2].text = f'{round(count/len(df["q13_overall_service_label"].dropna())*100, 1)}%'
    doc.add_page_break()
    
    # Section 5: Service Area Analysis
    doc.add_heading('5. Service Area Analysis', level=1)
    doc.add_heading('5.1 Mean Scores by Service Area', level=2)
    stats = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        data = df[col].dropna()
        if len(data) > 0:
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            stats.append({
                'Service Area': label,
                'Mean': round(data.mean(), 2),
                'Median': round(data.median(), 2),
                'Std Dev': round(data.std(), 2),
                'N': len(data)
            })
    
    if stats:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        headers = ['Service Area', 'Mean', 'Median', 'Std Dev', 'N']
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for s in sorted(stats, key=lambda x: x['Mean'], reverse=True):
            row = table.add_row().cells
            row[0].text = s['Service Area']
            row[1].text = f"{s['Mean']:.2f}"
            row[2].text = f"{s['Median']:.1f}"
            row[3].text = f"{s['Std Dev']:.2f}"
            row[4].text = str(s['N'])
    doc.add_page_break()
    
    # Section 6: Loyalty Indicators
    doc.add_heading('6. Loyalty Indicators', level=1)
    for col, title in [('q14_return_willingness_label', 'Willingness to Return'),
                       ('q15_recommend_label', 'Willingness to Recommend')]:
        if col in df.columns:
            idx = ['q14_return_willingness_label', 'q15_recommend_label'].index(col)
            doc.add_heading(f'6.{["a","b"][idx]} {title}', level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Response'; hdr[1].text = 'Count'; hdr[2].text = 'Percent'
            for response, count in df[col].value_counts().items():
                if not pd.isna(response):
                    row = table.add_row().cells
                    row[0].text = str(response)
                    row[1].text = str(count)
                    row[2].text = f'{round(count/len(df[col].dropna())*100, 1)}%'
    doc.add_page_break()
    
    # Section 7: Recommendations
    doc.add_heading('7. Recommendations', level=1)
    if stats:
        stats_df = pd.DataFrame(stats).sort_values('Mean', ascending=False)
        top_areas = stats_df.head(3)
        
        recommendations = f"""
Based on the survey findings, the following recommendations are proposed:

1. IMPROVE WATER AND BATHROOM FACILITIES (Priority #1)
   • This area received the lowest satisfaction score
   • Consider upgrading plumbing infrastructure
   • Ensure consistent hot water supply
   • Improve water pressure and bathroom cleanliness

2. ENHANCE Wi-Fi AND COMMUNICATION SERVICES (Priority #2)
   • Second lowest satisfaction area
   • Invest in better internet infrastructure
   • Ensure reliable connectivity throughout the facility

3. IMPROVE PROBLEM RESOLUTION SPEED
   • Guests reported delays in getting issues addressed
   • Implement a faster response system
   • Train staff on urgent issue handling

4. ADDRESS ELECTRICITY AND BASIC SERVICES
   • Ensure consistent power supply
   • Install backup generators for outages
   • Regular maintenance of electrical systems

5. MAINTAIN STRENGTHS (High-Performing Areas)
   • {top_areas.iloc[0]['Service Area'] if len(top_areas) > 0 else 'Environmental Safety'} scored highest
   • Continue current practices in high-performing areas
   • Maintain staff training programs
   • Implement regular guest feedback to monitor ongoing satisfaction
"""
        doc.add_paragraph(recommendations.strip())
    
    doc.save('survey_report_2018.docx')
    print("  ✅ Word report complete!")

# ============================================================
# WEB DASHBOARD
# ============================================================

def generate_dashboard(df):
    """Generate complete HTML dashboard with Raya University branding"""
    print("  Generating dashboard.html...")
    
    # Calculate metrics
    n = len(df)
    all_means = [df[col].mean() for col in df.columns if col.endswith('_num') and not pd.isna(df[col].mean())]
    overall_mean = np.mean(all_means) if all_means else 0
    
    means_data = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        m = df[col].mean()
        if not pd.isna(m):
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            means_data.append({'label': label, 'mean': m})
    means_df = pd.DataFrame(means_data).sort_values('mean', ascending=False)
    top_3 = means_df.head(3)['label'].tolist()
    
    return_rate = round(df['q14_return_willingness_label'].value_counts().get('Yes', 0) / len(df['q14_return_willingness_label'].dropna()) * 100, 1) if 'q14_return_willingness_label' in df.columns else 0
    rec_rate = round(df['q15_recommend_label'].value_counts().get('Yes', 0) / len(df['q15_recommend_label'].dropna()) * 100, 1) if 'q15_recommend_label' in df.columns else 0
    
    # Build HTML
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{UNIVERSITY_NAME} - Guest Satisfaction Survey Report 2018</title>
    <style>
        :root {{
            --navy: #1A2A5E;
            --gold: #C89B3C;
            --light-blue: #4A7FC1;
            --cream: #F5F0E6;
            --green: #63BE7B;
            --light-green: #A9D18E;
            --yellow: #FFEB84;
            --orange: #F4B183;
            --red: #E06666;
            --grey: #595959;
            --light-grey: #f0f2f5;
        }}
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
            background: var(--light-grey); 
            color: #333;
            line-height: 1.6;
        }}
        .container {{ max-width: 1300px; margin: 0 auto; padding: 20px; }}
        
        .header {{
            background: linear-gradient(135deg, var(--navy), #2a4a7a);
            color: white;
            padding: 35px 40px;
            border-radius: 12px;
            margin-bottom: 30px;
            text-align: center;
            border-bottom: 5px solid var(--gold);
        }}
        .header .university {{ font-size: 14px; opacity: 0.8; letter-spacing: 2px; }}
        .header h1 {{ font-size: 30px; font-weight: 700; margin: 5px 0; }}
        .header .department {{ font-size: 18px; opacity: 0.9; }}
        .header .subtitle {{ font-size: 16px; opacity: 0.85; margin-top: 5px; }}
        .header .date {{ font-size: 13px; opacity: 0.7; margin-top: 8px; }}
        
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 15px;
            margin-bottom: 30px;
        }}
        .stat-card {{
            background: white;
            padding: 20px 15px;
            border-radius: 10px;
            text-align: center;
            box-shadow: 0 2px 8px rgba(0,0,0,0.06);
            border-top: 4px solid var(--navy);
        }}
        .stat-card .number {{ font-size: 34px; font-weight: 700; color: var(--navy); }}
        .stat-card .label {{ font-size: 12px; color: #888; margin-top: 3px; }}
        .stat-card .number.green {{ color: var(--green); }}
        .stat-card .number.orange {{ color: var(--orange); }}
        .stat-card .number.gold {{ color: var(--gold); }}
        
        .section {{
            background: white;
            border-radius: 12px;
            padding: 25px 30px;
            margin-bottom: 25px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.06);
        }}
        .section h2 {{
            font-size: 22px;
            color: var(--navy);
            border-bottom: 3px solid var(--gold);
            padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        .section h3 {{
            font-size: 17px;
            color: var(--navy);
            margin-top: 20px;
            margin-bottom: 12px;
        }}
        .section h4 {{
            font-size: 14px;
            color: var(--grey);
            margin-top: 15px;
            margin-bottom: 8px;
        }}
        
        .chart-grid {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
        }}
        .chart-grid .full-width {{ grid-column: 1 / -1; }}
        .chart-card {{
            background: #fafbfc;
            border-radius: 8px;
            padding: 15px;
            border: 1px solid #e8ecf0;
        }}
        .chart-card img {{ width: 100%; height: auto; border-radius: 4px; }}
        .chart-card .caption {{
            font-size: 12px;
            color: #888;
            text-align: center;
            margin-top: 8px;
        }}
        
        .table-wrapper {{
            overflow-x: auto;
            margin: 10px 0 15px 0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 13px;
        }}
        table thead th {{
            background: var(--navy);
            color: white;
            padding: 8px 12px;
            text-align: left;
            font-weight: 600;
        }}
        table tbody td {{
            padding: 7px 12px;
            border-bottom: 1px solid #e8ecf0;
        }}
        table tbody tr:hover {{ background: #f5f7fa; }}
        table tbody tr:nth-child(even) {{ background: #fafbfc; }}
        table tbody tr:nth-child(even):hover {{ background: #f0f2f5; }}
        
        .badge {{
            display: inline-block;
            padding: 2px 10px;
            border-radius: 12px;
            font-size: 11px;
            font-weight: 600;
        }}
        .badge-green {{ background: var(--green); color: white; }}
        .badge-blue {{ background: var(--light-blue); color: white; }}
        .badge-orange {{ background: var(--orange); color: #333; }}
        .badge-red {{ background: var(--red); color: white; }}
        
        .report-text {{
            font-size: 14px;
            line-height: 1.8;
            color: #444;
        }}
        .report-text ul {{ padding-left: 25px; margin: 8px 0; }}
        .report-text ul li {{ margin-bottom: 4px; }}
        .highlight-box {{
            background: #f0f7ff;
            border-left: 4px solid var(--navy);
            padding: 15px 20px;
            border-radius: 4px;
            margin: 10px 0;
        }}
        .highlight-box.warning {{ background: #fff5f5; border-left-color: var(--red); }}
        
        .footer {{
            text-align: center;
            color: #aaa;
            font-size: 12px;
            padding: 25px 0 10px 0;
            border-top: 1px solid #e8ecf0;
        }}
        
        .tabs {{
            display: flex;
            flex-wrap: wrap;
            gap: 6px;
            margin: 10px 0 15px 0;
        }}
        .tab-btn {{
            padding: 6px 14px;
            border: 1px solid #d0d5dd;
            border-radius: 20px;
            background: white;
            cursor: pointer;
            font-size: 12px;
            transition: all 0.2s;
        }}
        .tab-btn:hover {{ background: #f0f2f5; }}
        .tab-btn.active {{
            background: var(--navy);
            color: white;
            border-color: var(--navy);
        }}
        .tab-content {{ display: none; }}
        .tab-content.active {{ display: block; }}
        
        .flex-between {{
            display: flex;
            justify-content: space-between;
            align-items: stretch;
            flex-wrap: wrap;
            gap: 15px;
        }}
        .flex-between .highlight-box {{ flex: 1; min-width: 200px; }}
        
        @media (max-width: 768px) {{
            .chart-grid {{ grid-template-columns: 1fr; }}
            .header h1 {{ font-size: 22px; }}
            .section {{ padding: 15px; }}
            .stats-grid {{ grid-template-columns: repeat(2, 1fr); }}
            .flex-between {{ flex-direction: column; }}
        }}
        @media print {{
            .header {{ background: var(--navy) !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            table thead th {{ background: var(--navy) !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
            .stat-card {{ border: 1px solid #ddd; }}
            .tab-content {{ display: block !important; }}
            .tab-btn {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        
        <div class="header">
            <div class="university">🏛️ {UNIVERSITY_NAME}</div>
            <h1>Guest Satisfaction Survey Report</h1>
            <div class="department">{DEPARTMENT}</div>
            <div class="subtitle">Survey Year: {SURVEY_YEAR}</div>
            <div class="date">Report Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}</div>
        </div>
        
        <div class="stats-grid">
            <div class="stat-card">
                <div class="number green">{overall_mean:.2f}</div>
                <div class="label">Overall Satisfaction (out of 5)</div>
            </div>
            <div class="stat-card">
                <div class="number gold">{return_rate}%</div>
                <div class="label">Would Return</div>
            </div>
            <div class="stat-card">
                <div class="number gold">{rec_rate}%</div>
                <div class="label">Would Recommend</div>
            </div>
            <div class="stat-card">
                <div class="number">{n}</div>
                <div class="label">Total Respondents</div>
            </div>
            <div class="stat-card">
                <div class="number orange">{len([c for c in df.columns if c.endswith('_num')])}</div>
                <div class="label">Service Areas Evaluated</div>
            </div>
        </div>
        
        <div class="section">
            <h2>1. Executive Summary</h2>
            <div class="report-text">
                <p>A guest satisfaction survey was conducted at the {DEPARTMENT} to assess the quality of service provided to visitors. A total of <strong>{n}</strong> guests completed the survey, rating their experience across reception, accommodation, support services, and overall satisfaction.</p>
                <br>
                <p>Overall, guests reported a high level of satisfaction, with an average score of <strong>{overall_mean:.2f} out of 5</strong> across all measured service dimensions. The strongest-performing areas were:</p>
                <ul>
'''
    for area in top_3[:3]:
        html += f'                    <li><strong>{area}</strong></li>\n'
    
    html += f'''                </ul>
                <p>while <strong>water and bathroom service</strong> received the most varied feedback and represents the clearest opportunity for improvement.</p>
                <br>
                <p><strong>{return_rate}%</strong> of respondents indicated they would return to the guesthouse, and <strong>{rec_rate}%</strong> said they would recommend it to others — both strong indicators of guest loyalty and satisfaction.</p>
            </div>
        </div>
        
        <div class="section">
            <h2>2. Methodology</h2>
            <div class="report-text">
                <p>Data was collected through a structured, bilingual (Amharic/English) paper-based questionnaire administered to guests of the {DEPARTMENT}. Responses were later digitized and analyzed.</p>
                <br>
                <ul>
                    <li><strong>Sample size:</strong> {n} respondents</li>
                    <li><strong>Survey instrument:</strong> 18-item structured questionnaire across six sections</li>
                    <li><strong>Rating scale:</strong> 5-point Likert scale (1 = Very dissatisfied/Very low, 5 = Very satisfied/Very good)</li>
                    <li><strong>Data collection method:</strong> Manual paper survey, later digitized for analysis</li>
                </ul>
                <br>
                <div class="highlight-box">
                    <strong>Note on scale:</strong> Questions 1-12 use satisfaction scale (ሀ=Very Satisfied → ሠ=Very Dissatisfied).<br>
                    Questions 13 and 16 use REVERSED quality scale (ሀ=Very Low → ሠ=Very Good).
                </div>
            </div>
        </div>
        
        <div class="section">
            <h2>3. Respondent Profile</h2>
            
            <div class="chart-grid">
                <div class="chart-card full-width">
                    <img src="charts/chart_visit_profile.png" alt="Visit Profile">
                    <div class="caption">Purpose of Visit and Length of Stay</div>
                </div>
            </div>
            
            <h3>3.1 Purpose of Visit</h3>
            <div class="table-wrapper">
                <table>
                    <thead><tr><th>Purpose</th><th>Count</th><th>Percent</th></tr></thead>
                    <tbody>
'''
    if 'visit_purpose_label' in df.columns:
        for purpose, count in df['visit_purpose_label'].value_counts().items():
            if not pd.isna(purpose):
                pct = round(count/len(df)*100, 1)
                html += f'                        <tr><td>{purpose}</td><td>{count}</td><td>{pct}%</td></tr>\n'
    
    html += f'''                    </tbody>
                </table>
            </div>
            
            <h3>3.2 Length of Stay</h3>
            <div class="table-wrapper">
                <table>
                    <thead><tr><th>Duration</th><th>Count</th><th>Percent</th></tr></thead>
                    <tbody>
'''
    if 'stay_duration_label' in df.columns:
        for duration, count in df['stay_duration_label'].value_counts().items():
            if not pd.isna(duration):
                pct = round(count/len(df)*100, 1)
                html += f'                        <tr><td>{duration}</td><td>{count}</td><td>{pct}%</td></tr>\n'
    
    html += f'''                    </tbody>
                </table>
            </div>
        </div>
        
        <div class="section">
            <h2>4. Overall Satisfaction</h2>
'''
    if 'q13_overall_service_num' in df.columns:
        overall_score = df['q13_overall_service_num'].mean()
        html += f'''
            <div class="report-text">
                <p>The overall service satisfaction score averages <strong>{overall_score:.2f} out of 5.0</strong>.</p>
            </div>
            
            <h3>4.1 Distribution of Overall Service Ratings</h3>
            <div class="table-wrapper">
                <table>
                    <thead><tr><th>Rating</th><th>Count</th><th>Percent</th></tr></thead>
                    <tbody>
'''
        valid_total = len(df['q13_overall_service_label'].dropna()) or 1
        for rating in ['Very Good', 'Good', 'Medium', 'Low', 'Very Low']:
            count = df['q13_overall_service_label'].value_counts().get(rating, 0)
            if count > 0:
                pct = round(count/valid_total*100, 1)
                html += f'                        <tr><td>{rating}</td><td>{count}</td><td>{pct}%</td></tr>\n'
        
        html += f'''                    </tbody>
                </table>
            </div>
'''
    
    html += f'''
        <div class="section">
            <h2>5. Visualizations</h2>
            
            <div class="chart-grid">
                <div class="chart-card full-width">
                    <img src="charts/chart_means.png" alt="Mean Scores">
                    <div class="caption">Mean Satisfaction Scores by Service Area</div>
                </div>
                <div class="chart-card full-width">
                    <img src="charts/chart_distribution.png" alt="Distribution">
                    <div class="caption">Distribution of Ratings for Key Service Areas</div>
                </div>
                <div class="chart-card full-width">
                    <img src="charts/chart_loyalty.png" alt="Loyalty">
                    <div class="caption">Loyalty Indicators</div>
                </div>
            </div>
        </div>
        
        <div class="section">
            <h2>6. Service Area Analysis</h2>
            
            <h3>6.1 Summary Statistics</h3>
            <div class="table-wrapper">
                <table>
                    <thead><tr><th>Service Area</th><th>N</th><th>Mean</th><th>Median</th><th>Std Dev</th><th>Min</th><th>Max</th></tr></thead>
                    <tbody>
'''
    
    # Generate stats table
    stats = []
    for col in [c for c in df.columns if c.endswith('_num')]:
        data = df[col].dropna()
        if len(data) > 0:
            label = QUESTION_LABELS.get(col.replace('_num', ''), col)
            stats.append({
                'Service Area': label,
                'N': len(data),
                'Mean': round(data.mean(), 2),
                'Median': round(data.median(), 2),
                'Std Dev': round(data.std(), 2),
                'Min': data.min(),
                'Max': data.max()
            })
    stats_df = pd.DataFrame(stats).sort_values('Mean', ascending=False)
    
    for _, row in stats_df.iterrows():
        badge = ''
        if row['Mean'] >= 4.5:
            badge = ' <span class="badge badge-green">High</span>'
        elif row['Mean'] >= 4.0:
            badge = ' <span class="badge badge-blue">Good</span>'
        elif row['Mean'] >= 3.5:
            badge = ' <span class="badge badge-orange">Medium</span>'
        else:
            badge = ' <span class="badge badge-red">Low</span>'
        html += f'                        <tr><td>{row["Service Area"]}{badge}</td><td>{row["N"]}</td><td>{row["Mean"]:.2f}</td><td>{row["Median"]:.1f}</td><td>{row["Std Dev"]:.2f}</td><td>{row["Min"]:.0f}</td><td>{row["Max"]:.0f}</td></tr>\n'
    
    html += f'''                    </tbody>
                </table>
            </div>
            
            <div class="flex-between">
                <div class="highlight-box">
                    <strong>🏆 Top Performing Areas:</strong><br>
'''
    for _, row in stats_df.head(3).iterrows():
        html += f'                    • {row["Service Area"]}: {row["Mean"]:.2f}/5.0<br>\n'
    
    html += f'''                </div>
                <div class="highlight-box warning">
                    <strong>⚠️ Areas for Improvement:</strong><br>
'''
    for _, row in stats_df.tail(3).iterrows():
        html += f'                    • {row["Service Area"]}: {row["Mean"]:.2f}/5.0<br>\n'
    
    html += f'''                </div>
            </div>
        </div>
        
        <div class="section">
            <h2>7. Loyalty Indicators</h2>
            
            <div class="chart-grid">
                <div class="chart-card full-width">
                    <img src="charts/chart_loyalty.png" alt="Loyalty">
                    <div class="caption">Return Intent and Recommendation Rates</div>
                </div>
            </div>
'''
    
    # Loyalty data
    for col, title in [('q14_return_willingness_label', 'Willingness to Return'),
                       ('q15_recommend_label', 'Willingness to Recommend')]:
        if col in df.columns:
            html += f'''
            <h3>7.{["a","b"][list(['q14_return_willingness_label','q15_recommend_label']).index(col)]} {title}</h3>
            <div class="table-wrapper">
                <table>
                    <thead><tr><th>Response</th><th>Count</th><th>Percent</th><th>Valid %</th></tr></thead>
                    <tbody>
'''
            for resp, count in df[col].value_counts(dropna=False).items():
                resp_label = resp if not pd.isna(resp) else 'Missing'
                pct = round(count/len(df)*100, 1)
                valid_pct = round(count/len(df[col].dropna())*100, 1) if len(df[col].dropna()) > 0 else 0
                html += f'                        <tr><td>{resp_label}</td><td>{count}</td><td>{pct}%</td><td>{valid_pct}%</td></tr>\n'
            html += f'''                    </tbody>
                </table>
            </div>
'''
    
    html += f'''
        </div>
        
        <div class="section">
            <h2>8. Recommendations</h2>
            <div class="report-text">
                <p>Based on the survey findings, the following recommendations are proposed:</p>
                <br>
                <h4>1. IMPROVE WATER AND BATHROOM FACILITIES (Priority #1)</h4>
                <ul>
                    <li>This area received the lowest satisfaction score</li>
                    <li>Consider upgrading plumbing infrastructure</li>
                    <li>Ensure consistent hot water supply</li>
                    <li>Improve water pressure and bathroom cleanliness</li>
                </ul>
                <br>
                <h4>2. ENHANCE Wi-Fi AND COMMUNICATION SERVICES (Priority #2)</h4>
                <ul>
                    <li>Second lowest satisfaction area</li>
                    <li>Invest in better internet infrastructure</li>
                    <li>Ensure reliable connectivity throughout the facility</li>
                </ul>
                <br>
                <h4>3. IMPROVE PROBLEM RESOLUTION SPEED</h4>
                <ul>
                    <li>Guests reported delays in getting issues addressed</li>
                    <li>Implement a faster response system</li>
                    <li>Train staff on urgent issue handling</li>
                </ul>
                <br>
                <h4>4. ADDRESS ELECTRICITY AND BASIC SERVICES</h4>
                <ul>
                    <li>Ensure consistent power supply</li>
                    <li>Install backup generators for outages</li>
                    <li>Regular maintenance of electrical systems</li>
                </ul>
                <br>
                <h4>5. MAINTAIN STRENGTHS (High-Performing Areas)</h4>
                <ul>
'''
    for area in top_3[:3]:
        html += f'                    <li>Continue current practices in <strong>{area}</strong></li>\n'
    html += f'''                    <li>Maintain staff training programs</li>
                    <li>Implement regular guest feedback to monitor ongoing satisfaction</li>
                </ul>
            </div>
        </div>
        
        <div class="footer">
            <p>{UNIVERSITY_NAME} &bull; {DEPARTMENT} &bull; Guest Satisfaction Survey {SURVEY_YEAR}</p>
            <p>Generated: {datetime.now().strftime("%B %d, %Y")}</p>
        </div>
        
    </div>
</body>
</html>'''
    
    with open('dashboard.html', 'w', encoding='utf-8') as f:
        f.write(html)
    print("  ✅ dashboard.html created!")

# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 60)
    print(f"{UNIVERSITY_NAME}")
    print(f"{DEPARTMENT}")
    print(f"Guest Satisfaction Survey Analysis {SURVEY_YEAR}")
    print("=" * 60)
    print("\nNote: Q13 and Q16 use REVERSED scale (ሀ=Very Low → ሠ=Very Good)")
    print("=" * 60)
    
    df = load_data()
    if df is None or len(df) == 0:
        print("\n❌ ERROR: No data loaded!")
        print("\nPlease make sure your Excel file is in this folder.")
        return
    
    df = process_data(df)
    
    print("\n📊 Generating outputs...")
    create_charts(df)
    export_excel(df)
    generate_word_report(df)
    generate_dashboard(df)
    
    print("\n" + "=" * 60)
    print("✅ ALL FILES CREATED SUCCESSFULLY!")
    print("=" * 60)
    print("\n📁 Output Files:")
    print("   🌐 dashboard.html           - Complete web dashboard (OPEN THIS!)")
    print("   📊 charts/chart_means.png   - Mean scores chart")
    print("   📊 charts/chart_visit_profile.png - Visit profile chart")
    print("   📊 charts/chart_loyalty.png - Loyalty indicators chart")
    print("   📊 charts/chart_distribution.png - Distribution chart")
    print("   📋 survey_analysis_2018.xlsx - SPSS-style tables")
    print("   📄 survey_report_2018.docx  - Word report")
    print("\n" + "=" * 60)
    print("🌐 OPEN dashboard.html IN YOUR BROWSER TO VIEW THE FULL REPORT")
    print("=" * 60)

if __name__ == "__main__":
    main()